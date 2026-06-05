# -*- coding: utf-8 -*-
"""Génère le rapport de projet Biblio Chat Bot en .docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

PRIMARY   = RGBColor(0x4D, 0x47, 0xA8)   # #4D47A8
DARK      = RGBColor(0x1A, 0x18, 0x25)   # #1A1825
MUTED     = RGBColor(0x8C, 0x89, 0xA0)   # #8C89A0
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xF3, 0xF2, 0xF8)   # #F3F2F8

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helper functions ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = PRIMARY
    run.font.name = 'Calibri'
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), '4D47A8')
    pb.append(bot)
    pPr.append(pb)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.name  = 'Calibri'
    run.font.color.rgb = DARK
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.name  = 'Calibri'
    run.font.color.rgb = DARK
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(1)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x4A, 0x46, 0x60)
    return p

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'E0DFF0')
    pb.append(bot)
    pPr.append(pb)

# ═════════════════════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ═════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('RAPPORT DE PROJET')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = MUTED
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('Biblio Chat Bot')
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = PRIMARY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Système de Gestion de Bibliothèque Intelligent\navec Chatbot IA Intégré')
run.font.size = Pt(14)
run.font.color.rgb = DARK
run.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

# Info table on cover
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

infos = [
    ('Réalisé par',   'Yasmine Ben Taher'),
    ('Filière',       '1ère Année Ingénierie'),
    ('Année',         '2025 – 2026'),
    ('Date',          datetime.date.today().strftime('%d/%m/%Y')),
]
for i, (label, value) in enumerate(infos):
    row = table.rows[i]
    set_cell_bg(row.cells[0], 'F3F2F8')
    lbl = row.cells[0].paragraphs[0]
    run = lbl.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = PRIMARY
    run.font.name = 'Calibri'
    val = row.cells[1].paragraphs[0]
    run2 = val.add_run(value)
    run2.font.size = Pt(10)
    run2.font.name = 'Calibri'

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  TABLE DES MATIÈRES (manuel)
# ═════════════════════════════════════════════════════════════════════════════

heading1('Table des Matières')
toc_items = [
    '1.  Contexte et Objectifs du Projet',
    '2.  Analyse des Besoins',
    '3.  Technologies Choisies et Justification',
    '4.  Architecture Globale de l\'Application',
    '5.  Conception de la Base de Données',
    '6.  Principales Fonctionnalités Développées',
    '7.  Intégration du Chatbot IA',
    '8.  Difficultés Rencontrées et Solutions',
    '9.  Conclusion et Perspectives d\'Amélioration',
]
for item in toc_items:
    bullet(item)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  1. CONTEXTE ET OBJECTIFS
# ═════════════════════════════════════════════════════════════════════════════

heading1('1. Contexte et Objectifs du Projet')

heading2('1.1 Contexte')
body(
    "Dans le cadre de notre formation en première année d'ingénierie, nous avons été amenés "
    "à réaliser un projet applicatif alliant développement web, programmation orientée objet "
    "et intelligence artificielle. Le projet Biblio Chat Bot répond à un besoin réel de "
    "modernisation de la gestion des bibliothèques, en proposant une solution numérique "
    "complète permettant de cataloguer, rechercher et recommander des ouvrages."
)
body(
    "Les bibliothèques traditionnelles font face à plusieurs défis : gestion manuelle des "
    "catalogues, difficulté à orienter les lecteurs vers des ouvrages pertinents, et absence "
    "d'interface intuitive pour les utilisateurs. Ce projet propose une réponse technologique "
    "à ces problématiques."
)

heading2('1.2 Objectifs')
body("Les objectifs principaux du projet sont :")
bullet("Développer un système de gestion de bibliothèque complet et fonctionnel")
bullet("Intégrer un chatbot IA local capable de recommander des livres et répondre aux questions des utilisateurs")
bullet("Proposer deux interfaces : une application web (Django) et une application bureau (CustomTkinter)")
bullet("Permettre l'import de catalogues de livres depuis des sources externes (Kaggle, Open Library)")
bullet("Assurer une expérience utilisateur moderne et intuitive avec une charte graphique cohérente")
bullet("Mettre en place un panneau d'administration complet pour la gestion des données")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  2. ANALYSE DES BESOINS
# ═════════════════════════════════════════════════════════════════════════════

heading1('2. Analyse des Besoins')

heading2('2.1 Besoins Fonctionnels')

heading2('Gestion du Catalogue')
bullet("Affichage de la liste des livres avec recherche et filtres (titre, auteur, catégorie, disponibilité)")
bullet("Fiche détaillée de chaque livre (description, auteur, note, pages, langue, ISBN, année)")
bullet("Un livre peut appartenir à plusieurs catégories simultanément (relation Many-to-Many)")
bullet("Gestion des disponibilités (exemplaires disponibles / total)")
bullet("Affichage automatique des couvertures de livres (Open Library API)")

heading2('Gestion des Utilisateurs et Administration')
bullet("Interface d'administration complète pour créer, modifier et supprimer des livres, auteurs et catégories")
bullet("Authentification sécurisée pour l'accès à l'administration")
bullet("Import de données depuis un fichier CSV (dataset Kaggle GoodReads)")
bullet("Commande de gestion pour télécharger automatiquement les couvertures")

heading2('Chatbot IA')
bullet("Chat en temps réel avec streaming des réponses (Server-Sent Events)")
bullet("Connaissance du catalogue complet de la bibliothèque (contexte injecté dans le prompt)")
bullet("Persistance de l'historique des conversations en base de données")
bullet("Support de plusieurs modèles de langage via Ollama")
bullet("Recommandations personnalisées basées sur les préférences de l'utilisateur")

heading2('2.2 Besoins Non Fonctionnels')
bullet("Performance : réponses de l'IA streamées en temps réel sans blocage de l'interface")
bullet("Utilisabilité : interface claire, intuitive et cohérente visuellement")
bullet("Maintenabilité : code structuré selon le pattern MVC de Django")
bullet("Extensibilité : architecture modulaire permettant l'ajout de nouvelles fonctionnalités")
bullet("Sécurité : protection des routes d'administration, validation des données")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  3. TECHNOLOGIES
# ═════════════════════════════════════════════════════════════════════════════

heading1('3. Technologies Choisies et Justification')

heading2('3.1 Backend Web — Django 6.0')
body(
    "Django a été choisi comme framework principal pour plusieurs raisons. Son ORM (Object-Relational "
    "Mapper) permet de manipuler la base de données via des classes Python sans écrire de SQL. "
    "Son système de templates facilite la génération de pages HTML dynamiques. L'interface "
    "d'administration auto-générée accélère considérablement le développement. Enfin, son "
    "architecture MVT (Modèle-Vue-Template) impose une séparation claire des responsabilités."
)

heading2('3.2 Interface Bureau — CustomTkinter')
body(
    "CustomTkinter est une extension moderne de la bibliothèque Tkinter standard de Python. "
    "Elle offre des widgets au design contemporain (coins arrondis, thèmes sombre/clair, "
    "animations), tout en restant légère et sans dépendance externe complexe. Elle a été "
    "choisie pour créer une application bureau qui réutilise directement l'ORM Django, "
    "évitant ainsi la duplication de la logique métier."
)

heading2('3.3 Intelligence Artificielle — Ollama + LLaMA 3.2')
body(
    "Ollama est une solution d'IA locale qui permet d'exécuter des modèles de langage "
    "(LLM) directement sur la machine, sans envoyer de données vers des serveurs externes. "
    "Le modèle LLaMA 3.2 (3.2 milliards de paramètres, quantifié Q4_K_M) offre un bon "
    "équilibre entre performance et qualité des réponses. L'API REST d'Ollama est simple "
    "et supporte le streaming natif."
)

heading2('3.4 Base de Données — SQLite')
body(
    "SQLite a été utilisé en phase de développement pour sa simplicité de configuration "
    "(aucun serveur requis, fichier unique db.sqlite3). La migration vers PostgreSQL est "
    "prévue pour la production, Django supportant les deux sans modification du code applicatif."
)

heading2('3.5 Frontend Web — Tailwind CSS')
body(
    "Tailwind CSS est un framework CSS utilitaire qui permet de styliser les composants "
    "directement dans le HTML via des classes prédéfinies. Il a été intégré via CDN pour "
    "accélérer le développement. Les polices Playfair Display et Inter (Google Fonts) "
    "apportent une identité visuelle élégante."
)

heading2('3.6 Open Library API')
body(
    "L'API gratuite d'Open Library (openlibrary.org) permet de rechercher des couvertures "
    "de livres par titre ou ISBN. Les images sont téléchargées et stockées localement dans "
    "le répertoire media/book_covers/, ce qui évite les requêtes répétées."
)

heading2('3.7 Jazzmin — Thème Admin')
body(
    "Jazzmin est un package Django qui remplace l'interface d'administration par défaut "
    "par une interface moderne basée sur AdminLTE 3 et Bootstrap 4. Il permet une "
    "personnalisation complète (couleurs, icônes, menus) via la configuration settings.py."
)

# Tableau récapitulatif
doc.add_paragraph()
heading2('Tableau Récapitulatif des Technologies')
table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'

headers = ['Technologie', 'Version', 'Rôle']
for j, h in enumerate(headers):
    cell = table.rows[0].cells[j]
    set_cell_bg(cell, '4D47A8')
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.name = 'Calibri'
    run.font.size = Pt(10)

rows_data = [
    ('Django',          '6.0',    'Framework web backend, ORM, admin'),
    ('CustomTkinter',   '5.2',    'Interface bureau Python'),
    ('Ollama + LLaMA',  '3.2B',   'Moteur IA local (chatbot)'),
    ('SQLite / PostgreSQL', '3 / 15', 'Base de données'),
    ('Tailwind CSS',    'CDN',    'Styles frontend web'),
    ('Open Library API','—',      'Couvertures de livres'),
    ('Jazzmin',         '3.0',    'Thème interface d\'administration'),
]
for i, (tech, ver, role) in enumerate(rows_data):
    row = table.rows[i + 1]
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, 'F3F2F8')
    row.cells[0].paragraphs[0].add_run(tech).font.name = 'Calibri'
    row.cells[1].paragraphs[0].add_run(ver).font.name  = 'Calibri'
    row.cells[2].paragraphs[0].add_run(role).font.name = 'Calibri'
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  4. ARCHITECTURE
# ═════════════════════════════════════════════════════════════════════════════

heading1('4. Architecture Globale de l\'Application')

heading2('4.1 Vue d\'Ensemble')
body(
    "L'application suit une architecture en couches inspirée du patron MVC (Modèle-Vue-Contrôleur), "
    "adapté au pattern MVT (Modèle-Vue-Template) de Django. Deux interfaces distinctes "
    "coexistent et partagent la même couche de données."
)

heading2('4.2 Structure des Applications Django')
body("Le projet est organisé en deux applications Django :")
bullet("library : gestion du catalogue (livres, auteurs, catégories), vues de navigation")
bullet("chat : gestion des conversations IA (modèles Conversation et Message, vues SSE)")

heading2('4.3 Flux de Données — Application Web')
code_block("Navigateur → URL Router (urls.py) → Vue Django (views.py)")
code_block("Vue → ORM Django → SQLite/PostgreSQL")
code_block("Vue → Template HTML + Tailwind CSS → Navigateur")
code_block("Chat : Vue → Ollama API (streaming SSE) → JavaScript → Interface")

heading2('4.4 Flux de Données — Application Bureau')
code_block("CustomTkinter UI → Django ORM (direct, sans HTTP)")
code_block("Thread principal : rendu UI")
code_block("Thread secondaire : requête Ollama API + mise à jour UI via after()")
code_block("Thread tertiaire : téléchargement couvertures Open Library")

heading2('4.5 Structure des Fichiers')
code_block("Biblio_Chat_Bot/")
code_block("├── Biblio_Chat_Bot/        # Configuration Django")
code_block("│   ├── settings.py         # Paramètres + Jazzmin")
code_block("│   └── urls.py             # Routage principal")
code_block("├── library/                # App catalogue")
code_block("│   ├── models.py           # Book, Author, Category")
code_block("│   ├── views.py            # Vues web")
code_block("│   ├── admin.py            # Interface admin")
code_block("│   └── management/commands/ # fetch_covers, import_books, seed_data")
code_block("├── chat/                   # App chatbot")
code_block("│   ├── models.py           # Conversation, Message")
code_block("│   └── views.py            # SSE streaming")
code_block("├── templates/              # Templates HTML")
code_block("├── static/                 # CSS custom admin")
code_block("├── media/book_covers/      # Couvertures téléchargées")
code_block("├── desktop_app.py          # Application CustomTkinter")
code_block("├── rose_theme.json         # Thème CustomTkinter personnalisé")
code_block("└── db.sqlite3              # Base de données")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  5. BASE DE DONNÉES
# ═════════════════════════════════════════════════════════════════════════════

heading1('5. Conception de la Base de Données')

heading2('5.1 Diagramme Entité-Association')
body(
    "La base de données comprend 5 tables principales issues de deux applications Django. "
    "Les relations sont gérées par l'ORM Django via des ForeignKey et ManyToManyField."
)

heading2('5.2 Modèle Category')
code_block("Category(id, name, description, icon)")
body("Représente un genre littéraire (Fiction, Science-Fiction, Philosophie, etc.) "
     "avec une icône emoji pour l'identification visuelle rapide.")

heading2('5.3 Modèle Author')
code_block("Author(id, name, bio)")
body("Stocke les informations biographiques des auteurs, reliés aux livres par une ForeignKey.")

heading2('5.4 Modèle Book')
code_block("Book(id, title, author[FK], categories[M2M], description, isbn,")
code_block("     published_year, available_copies, total_copies, cover_color,")
code_block("     cover_image, rating, pages, language, created_at)")
body(
    "Entité centrale du système. La relation ManyToMany avec Category (introduite lors d'une "
    "migration) permet à un livre d'appartenir à plusieurs genres simultanément. "
    "cover_image stocke le chemin vers l'image téléchargée depuis Open Library. "
    "cover_color est utilisé comme couleur de fond de substitution."
)

heading2('5.5 Modèle Conversation')
code_block("Conversation(id, session_key, created_at, updated_at)")
body("Regroupe les messages d'une session de chat. session_key = 'desktop_app' "
     "pour l'application bureau, ou la clé de session Django pour le web.")

heading2('5.6 Modèle Message')
code_block("Message(id, conversation[FK], role['user'|'assistant'], content, created_at)")
body("Stocke chaque échange entre l'utilisateur et l'IA. "
     "L'historique est rechargé à chaque ouverture du chat, assurant la continuité des conversations.")

heading2('5.7 Relations')

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
for j, h in enumerate(['Relation', 'Type', 'Description']):
    cell = table.rows[0].cells[j]
    set_cell_bg(cell, '4D47A8')
    run = cell.paragraphs[0].add_run(h)
    run.bold = True; run.font.color.rgb = WHITE
    run.font.name = 'Calibri'; run.font.size = Pt(10)

rels = [
    ('Book → Author',       'ForeignKey (N→1)',    'Un livre a un auteur principal'),
    ('Book ↔ Category',     'ManyToMany',          'Un livre peut avoir plusieurs catégories'),
    ('Message → Conversation','ForeignKey (N→1)', 'Un message appartient à une conversation'),
]
for i, (rel, typ, desc) in enumerate(rels):
    row = table.rows[i + 1]
    if i % 2 == 0:
        for cell in row.cells: set_cell_bg(cell, 'F3F2F8')
    for j, val in enumerate([rel, typ, desc]):
        r = row.cells[j].paragraphs[0].add_run(val)
        r.font.name = 'Calibri'; r.font.size = Pt(10)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  6. FONCTIONNALITÉS
# ═════════════════════════════════════════════════════════════════════════════

heading1('6. Principales Fonctionnalités Développées')

heading2('6.1 Catalogue de Livres (Web & Bureau)')
body(
    "La page principale du catalogue affiche l'ensemble des livres disponibles avec leurs "
    "couvertures, notes, auteurs et catégories. Un système de recherche full-text (titre, "
    "auteur, description) et des filtres par catégorie et disponibilité permettent une "
    "navigation efficace."
)
bullet("Vue liste avec barre de recherche, filtres et toggle grille/liste")
bullet("Vue détail avec couverture, description complète, biographie de l'auteur")
bullet("Livres « similaires » basés sur les catégories communes")
bullet("Bouton direct « Demander à Biblio AI » depuis la fiche livre")

heading2('6.2 Gestion des Catégories (ManyToMany)')
body(
    "Une évolution majeure du projet a été la migration du champ category (ForeignKey) "
    "vers categories (ManyToManyField). Cette migration préserve les données existantes "
    "via une RunPython dans la migration Django, et enrichit l'expérience en permettant "
    "des classifications croisées (ex : un roman peut être à la fois Fiction et Philosophie)."
)

heading2('6.3 Import de Données — Kaggle GoodReads')
body(
    "Une commande de gestion Django (import_books) permet d'importer des livres depuis "
    "le dataset GoodReads de Kaggle (11 127 livres au format CSV). La commande gère :"
)
bullet("La création ou la récupération des auteurs (get_or_create)")
bullet("Le déduplication par titre + auteur")
bullet("La normalisation des codes langue, des dates de publication et des notes")
bullet("Un paramètre --limit pour contrôler le volume importé")

heading2('6.4 Couvertures Automatiques — Open Library')
body(
    "La commande fetch_covers interroge l'API d'Open Library pour chaque livre sans "
    "couverture, télécharge l'image (format JPEG, résolution L) et la sauvegarde dans "
    "media/book_covers/ via le champ ImageField de Django. Les couvertures sont également "
    "mises en cache dans .cover_cache/ pour l'application bureau."
)

heading2('6.5 Interface Bureau — CustomTkinter')
body(
    "L'application bureau propose quatre écrans principaux accessibles depuis une barre de "
    "navigation latérale avec icônes et labels :"
)
bullet("Accueil : bannière compacte avec recherche intégrée, statistiques, recommandation vedette, ajouts récents, top notes")
bullet("Catalogue : vue grille (4 colonnes) ou liste avec filtres, switch « Disponibles uniquement »")
bullet("Catégories : grille de cartes avec icônes et compteurs de livres")
bullet("Chat IA : interface de messagerie avec historique persistant, sélection de modèle, suggestions")

heading2('6.6 Panneau d\'Administration — Jazzmin')
body(
    "L'interface d'administration Django a été complètement redessinée avec le package "
    "Jazzmin pour correspondre à la charte graphique du projet (couleurs #4D47A8, "
    "polices Playfair Display + Inter). Elle offre :"
)
bullet("Gestion CRUD complète des livres, auteurs et catégories")
bullet("Widget filter_horizontal pour l'assignation multiple de catégories")
bullet("Barre de recherche et filtres par catégorie, langue et disponibilité")
bullet("Lien direct vers l'interface bibliothèque depuis le menu")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  7. CHATBOT IA
# ═════════════════════════════════════════════════════════════════════════════

heading1('7. Intégration du Chatbot IA et son Fonctionnement')

heading2('7.1 Architecture de l\'IA')
body(
    "Le chatbot repose sur Ollama, un serveur d'inférence local qui expose une API REST "
    "compatible avec le standard OpenAI Chat Completions. Le modèle LLaMA 3.2 (3.2B paramètres "
    "quantifiés Q4_K_M) est exécuté entièrement en local, garantissant la confidentialité "
    "des données et l'absence de coûts d'utilisation."
)

heading2('7.2 Prompt Système')
body(
    "À chaque requête, un prompt système est construit dynamiquement en incluant le catalogue "
    "complet de la bibliothèque. Cette technique, appelée RAG simplifié (Retrieval-Augmented "
    "Generation), permet au modèle de répondre avec précision sur les livres disponibles :"
)
code_block("Vous êtes Biblio, un assistant bibliothécaire compétent et chaleureux.")
code_block("Aidez les utilisateurs à découvrir des livres et à obtenir des recommandations.")
code_block("")
code_block("Catalogue actuel de la bibliothèque :")
code_block("  • '1984' par George Orwell [Fiction, ★4.7] — disponible")
code_block("  • 'Méditations' par Marc Aurèle [Philosophie, ★4.9] — disponible")
code_block("  ...")

heading2('7.3 Streaming des Réponses')
body(
    "Les réponses de l'IA sont transmises en temps réel via deux mécanismes distincts :"
)
bullet("Application Web : Server-Sent Events (SSE) — le serveur Django maintient une connexion HTTP ouverte et envoie les tokens au fur et à mesure")
bullet("Application Bureau : Thread Python dédié qui consomme l'API /api/chat d'Ollama en mode stream=True et met à jour le label CTk via self.after(0, ...)")

heading2('7.4 Persistance des Conversations')
body(
    "Chaque message (rôle user ou assistant) est sauvegardé en base de données via le modèle "
    "Message. L'historique complet est rechargé à chaque ouverture de la page Chat, "
    "assurant la continuité des échanges entre les sessions. La fonction Clear supprime "
    "la conversation en base et en crée une nouvelle."
)

heading2('7.5 Sécurité Thread — Application Bureau')
body(
    "L'ORM Django utilise des connexions base de données par thread. Pour éviter les erreurs "
    "de connexions périmées dans les threads CustomTkinter, close_old_connections() est "
    "appelé au début de chaque fonction ORM exécutée hors du thread principal."
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  8. DIFFICULTÉS ET SOLUTIONS
# ═════════════════════════════════════════════════════════════════════════════

heading1('8. Difficultés Rencontrées et Solutions Apportées')

table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'

for j, h in enumerate(['Difficulté', 'Cause', 'Solution']):
    cell = table.rows[0].cells[j]
    set_cell_bg(cell, '4D47A8')
    run = cell.paragraphs[0].add_run(h)
    run.bold = True; run.font.color.rgb = WHITE
    run.font.name = 'Calibri'; run.font.size = Pt(10)

difficulties = [
    (
        'Migration ForeignKey → ManyToMany',
        'Django ne permet pas la conversion directe sans données intermédiaires',
        'Migration manuelle en 4 étapes : AddField (related_name temporaire), RunPython (copie données), RemoveField (ancienne FK), AlterField (renommage)'
    ),
    (
        'Couleurs 8 chiffres (#RRGGBBAA) dans tkinter',
        'tkinter ne supporte pas le canal alpha dans les codes hexadécimaux',
        'Remplacement par des couleurs 6 chiffres équivalentes sans transparence'
    ),
    (
        'SSL Certificate Error sur Windows',
        'Python 3.14 sur Windows ne dispose pas des certificats CA système dans son store SSL',
        'Ajout de verify=False dans toutes les requêtes requests + urllib3.disable_warnings()'
    ),
    (
        'ORM Django dans threads CustomTkinter',
        'Django ferme les connexions DB entre les cycles de requêtes web',
        'Appel de close_old_connections() en début de chaque fonction ORM exécutée en thread secondaire'
    ),
    (
        'Widgets CTk détruits pendant streaming',
        'L\'utilisateur peut naviguer pendant qu\'Ollama génère une réponse',
        'Méthode _try(widget, **kw) qui capture silencieusement les TclError, et réinitialisation du flag _streaming dans _clear_main()'
    ),
    (
        'Encodage emoji Windows terminal',
        'Le terminal Windows (cp1252) ne supporte pas les caractères Unicode emoji',
        'Remplacement des emojis dans les messages de management commands par du texte ASCII ([fetch], [skip], etc.)'
    ),
]

for i, (diff, cause, sol) in enumerate(difficulties):
    row = table.rows[i + 1]
    if i % 2 == 0:
        for cell in row.cells: set_cell_bg(cell, 'F3F2F8')
    for j, val in enumerate([diff, cause, sol]):
        r = row.cells[j].paragraphs[0].add_run(val)
        r.font.name = 'Calibri'; r.font.size = Pt(9)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  9. CONCLUSION
# ═════════════════════════════════════════════════════════════════════════════

heading1('9. Conclusion et Perspectives d\'Amélioration')

heading2('9.1 Bilan du Projet')
body(
    "Le projet Biblio Chat Bot constitue une réalisation complète et fonctionnelle d'un "
    "système de gestion de bibliothèque moderne. Il intègre avec succès une application "
    "web Django, une application bureau CustomTkinter, un chatbot IA local via Ollama, "
    "et un panneau d'administration professionnel. Le catalogue compte actuellement "
    "150 ouvrages avec couvertures, organisés par catégories multiples."
)
body(
    "Ce projet a permis de maîtriser des compétences techniques variées : architecture "
    "MVC/MVT, ORM Django, migrations de base de données, programmation multithreadée, "
    "intégration d'API REST, streaming de données en temps réel, et conception d'interfaces "
    "graphiques. Il illustre également l'importance de la gestion des erreurs, de la "
    "thread-safety et de la cohérence visuelle dans un projet multi-interfaces."
)

heading2('9.2 Perspectives d\'Amélioration')

bullet("Migration vers PostgreSQL : configurer les credentials et migrer les données pour une utilisation en production")
bullet("Authentification utilisateurs : système de comptes, emprunts personnalisés, historique de lecture")
bullet("Système d'emprunt complet : réservations, dates de retour, notifications")
bullet("Amélioration du RAG : vectorisation des descriptions avec ChromaDB ou Faiss pour des recommandations sémantiques plus précises")
bullet("Support multilingue : traduction de l'interface et indexation des livres en plusieurs langues")
bullet("Application mobile : API REST Django REST Framework + application React Native ou Flutter")
bullet("Déploiement : conteneurisation Docker, déploiement sur un serveur avec PostgreSQL et Nginx")
bullet("Tests automatisés : tests unitaires et d'intégration avec pytest-django")
bullet("Export de données : génération de rapports PDF, export Excel du catalogue")
bullet("Statistiques avancées : tableau de bord admin avec graphiques (livres les plus demandés, tendances)")

heading2('9.3 Compétences Acquises')
body("Ce projet a permis de développer et consolider les compétences suivantes :")
bullet("Développement web fullstack avec Django (ORM, migrations, templates, admin)")
bullet("Programmation GUI avec Python (CustomTkinter, threading, events)")
bullet("Intégration d'IA locale (Ollama, LLM, prompt engineering, streaming SSE)")
bullet("Conception et migration de base de données relationnelle")
bullet("Gestion de projet : versioning Git, débogage, documentation")

divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('— Fin du rapport —')
run.font.color.rgb = MUTED
run.font.size = Pt(11)
run.font.name = 'Calibri'
run.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
output = r'C:\Users\Yasmine BEN TAHER\Desktop\Yasmine_1ING\Biblio_Chat_Bot\Rapport_Biblio_Chat_Bot.docx'
doc.save(output)
print(f'Rapport genere : {output}')
